from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io

def generate_report(df, value_col, factor_col, result):
    doc = Document()
    doc.add_heading(f"Аналіз показника '{value_col}'", level=1)
    doc.add_paragraph(f"Фактор: {factor_col}")
    doc.add_paragraph(" ")

    doc.add_heading("Початкові дані", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_heading("Результати аналізу", level=2)
    table2 = doc.add_table(rows=1, cols=len(result["table"].columns))
    hdr2 = table2.rows[0].cells
    for i, col in enumerate(result["table"].columns):
        hdr2[i].text = col
    for _, row in result["table"].iterrows():
        row_cells = table2.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_heading("Сила впливу факторів (η²)", level=2)
    table3 = doc.add_table(rows=1, cols=2)
    table3.rows[0].cells[0].text = "Фактор"
    table3.rows[0].cells[1].text = "η²"
    for _, row in result["effect_size"].iterrows():
        row_cells = table3.add_row().cells
        row_cells[0].text = str(row["Фактор"])
        row_cells[1].text = f"{row['η²']:.4f}"

    fig = plt.figure()
    ax = fig.add_subplot()
    df.boxplot(column=value_col, by=factor_col, ax=ax)
    plt.tight_layout()
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png')
    img_stream.seek(0)
    doc.add_heading("Графік розподілу", level=2)
    doc.add_picture(img_stream, width=Inches(5.5))

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream
